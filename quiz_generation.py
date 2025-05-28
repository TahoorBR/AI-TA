import requests
from docx import Document
import os
from flask import send_from_directory

# Replace with your actual API key
API_KEY = "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"
API_URL = "https://openrouter.ai/api/v1/chat/completions"

def generate_quiz(quiz_data):
    # Formulate prompt based on quiz data
    prompt = (
        f"Create a quiz on the topic '{quiz_data['topic']}' for {quiz_data['education_level']} level students.\n"
        f"Quiz Type: {quiz_data['quiz_type']}\nNumber of Questions: {quiz_data['num_questions']}\n\n"
        "Please list all questions first and provide the answers at the end in the following format:\n"
        "Questions:\n"
    )

    if quiz_data['quiz_type'] == "mcqs":
        prompt += "1. Question [MCQ options: A, B, C, D]\n"
    elif quiz_data['quiz_type'] == "true/false":
        prompt += "1. Question [Answer: True or False]\n"
    elif quiz_data['quiz_type'] == "shortquestions":
        prompt += "1. Question [Short Answer]\n"
    elif quiz_data['quiz_type'] == "longquestions":
        prompt += "1. Question [Long Answer]\n"
    prompt += "\nAnswers:\n1. Answer\n"

    # Call OpenRouter API with LLaMA 3.2 3B Instruct model
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost",  # Replace with your website or repo if you publish
        "X-Title": "AI Quiz Generator"
    }

    payload = {
        "model": "meta-llama/llama-3.2-3b-instruct:free",
        "messages": [
            {"role": "system", "content": "You are an educational assistant that generates quizzes."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }

    try:
        response = requests.post(API_URL, headers=headers, json=payload)
        response.raise_for_status()
        quiz_text = response.json()['choices'][0]['message']['content']
    except Exception as e:
        return f"Error generating quiz: {e}"

    # Save the quiz to a DOCX file
    doc = Document()
    doc.add_heading("AI Teaching Assistant Quiz", level=1)

    for line in quiz_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    output_dir = os.path.join("static", "quizzes")
    os.makedirs(output_dir, exist_ok=True)
    doc.save(os.path.join(output_dir, "quiz.docx"))

    return quiz_text

# Route for downloading the quiz
def download_quiz():
    return send_from_directory(
        os.path.join("static", "quizzes"),
        "quiz.docx",
        as_attachment=True
    )
