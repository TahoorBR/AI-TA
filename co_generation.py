import requests
from docx import Document
import os
from flask import send_from_directory

# Your OpenRouter API Key
API_KEY = "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"
API_URL = "https://openrouter.ai/api/v1/chat/completions"

def generate_course_outline(course_data):
    """
    Generate a course outline document using LLaMA 3.2 3B via OpenRouter,
    based on minimal user input and HEC format guidelines.
    """
    prompt = (
        f"Create a detailed course outline following HEC guidelines for the course '{course_data['course_title']}' "
        f"({course_data['course_code']}) at {course_data['education_level']} level.\n\n"
        f"Only the following data is provided:\n"
        f"Course Title: {course_data['course_title']}\n"
        f"Course Code: {course_data['course_code']}\n"
        f"Education Level: {course_data['education_level']}\n"
        f"Credit Hours: {course_data['credit_hours']}\n"
        f"Prerequisites: {course_data.get('prerequisites', 'None')}\n\n"
        f"Based on this, generate the following sections:\n"
        f"1. Course Description\n"
        f"2. Course Objectives (list)\n"
        f"3. Intended Learning Outcomes (list)\n"
        f"4. Teaching and Learning Methods (list)\n"
        f"5. Assessment and Evaluation Methods with weightage (list)\n"
        f"6. Weekly Course Outline (week number and topics)\n"
        f"7. Recommended Textbooks and References (list)\n\n"
        f"Format the output clearly with headings and bullet points suitable for a Word document."
    )

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost",  # Replace with your app domain if public
        "X-Title": "Course Outline Generator"
    }

    payload = {
        "model": "meta-llama/llama-3.2-3b-instruct:free",
        "messages": [
            {"role": "system", "content": "You are a helpful assistant that writes course outlines based on HEC format."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }

    try:
        response = requests.post(API_URL, headers=headers, json=payload)
        response.raise_for_status()
        course_outline_text = response.json()['choices'][0]['message']['content']
    except Exception as e:
        return f"Error generating course outline: {e}"

    # Create Word document
    doc = Document()
    doc.add_heading("Course Outline (HEC Guidelines)", level=1)
    doc.add_paragraph(f"Course Title: {course_data['course_title']}")
    doc.add_paragraph(f"Course Code: {course_data['course_code']}")
    doc.add_paragraph(f"Education Level: {course_data['education_level']}")
    doc.add_paragraph(f"Credit Hours: {course_data['credit_hours']}")
    doc.add_paragraph(f"Prerequisites: {course_data.get('prerequisites', 'None')}")
    doc.add_paragraph("")  # Blank line

    for line in course_outline_text.strip().split('\n'):
        line = line.strip()
        if line.endswith(":") and len(line) < 40:
            doc.add_heading(line[:-1], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='ListBullet')
        else:
            doc.add_paragraph(line)

    output_dir = os.path.join("static", "course_outlines")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, "course_outline.docx")
    doc.save(doc_path)

    return course_outline_text

def download_course_outline():
    """
    Serve the generated course outline DOCX file for download.
    """
    return send_from_directory(
        os.path.join("static", "course_outlines"),
        "course_outline.docx",
        as_attachment=True
    )
