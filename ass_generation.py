import requests
from docx import Document
import os
from flask import send_from_directory

# OpenRouter API config
API_KEY = "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"
API_URL = "https://openrouter.ai/api/v1/chat/completions"


def generate_assignment(assignment_data):
    # Extract data
    topic = assignment_data['topic']
    education_level = assignment_data['education_level']
    assignment_type = assignment_data['assignment_type'].lower()
    num_questions = assignment_data['num_questions']
    word_limit = assignment_data.get('word_limit', None)

    # Build prompt
    prompt = (
        f"Create an assignment on the topic '{topic}' for {education_level} level students.\n"
        f"Assignment Type: {assignment_data['assignment_type']}\n"
        f"Number of Questions: {num_questions}\n"
    )

    if assignment_type in ['long questions', 'essay', 'research'] and word_limit:
        prompt += f"Each answer should be approximately {word_limit} words.\n"

    prompt += (
        "\nPlease list all questions first and provide the answers at the end in the following format:\n"
        "Questions:\n1. Question\n...\n\nAnswers:\n1. Answer\n...\n"
    )

    # Prepare the OpenRouter API payload
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "meta-llama/llama-3.2-3b-instruct:free",
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    # Call the OpenRouter API
    response = requests.post(API_URL, headers=headers, json=payload)
    response.raise_for_status()
    assignment_text = response.json()["choices"][0]["message"]["content"]

    # Save to DOCX
    doc = Document()
    doc.add_heading("AI Teaching Assistant Assignment", level=1)

    for line in assignment_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    output_dir = os.path.join("static", "assignments")
    os.makedirs(output_dir, exist_ok=True)
    doc.save(os.path.join(output_dir, "assignment.docx"))

    return assignment_text


def download_assignment():
    return send_from_directory(
        os.path.join("static", "assignments"),
        "assignment.docx",
        as_attachment=True
    )
