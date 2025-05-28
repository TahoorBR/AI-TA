import os
import requests
from docx import Document
from flask import send_from_directory

# OCR-related imports
import pytesseract
from PIL import Image
import cv2

# OpenRouter API config
API_KEY = "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"
API_URL = "https://openrouter.ai/api/v1/chat/completions"


def extract_text_from_image(image_path):
    """Extracts text from an image using Tesseract OCR."""
    image = cv2.imread(image_path)
    if image is None:
        raise ValueError(f"Could not read image: {image_path}")
    image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    text = pytesseract.image_to_string(Image.fromarray(image))
    return text.strip()


def generate_feedback(quiz_data, image_path=None):
    subject = quiz_data.get("subject", "General")
    education_level = quiz_data.get("education_level", "High School")
    assignment_type = quiz_data.get("assignment_type", "Assignment")
    total_marks = quiz_data.get("total_marks", 100)

    # Determine the submission text
    if image_path:
        try:
            submission_text = extract_text_from_image(image_path)
        except Exception as e:
            return f"❌ OCR Failed: {str(e)}"
    else:
        submission_text = quiz_data.get("submission_text", "")

    # Compose the grading prompt
    prompt = (
        f"You are an expert {subject} teacher assessing a {assignment_type} submitted by a {education_level} level student.\n\n"
        f"The total marks for this assignment are {total_marks}. All scores must be strictly out of {total_marks}, not 100.\n\n"
        "Your task:\n"
        "1. Create an appropriate grading rubric tailored to this type of work.\n"
        "2. Use the rubric to evaluate the submission. Your marking must be extremely strict:\n"
        "   - First, independently generate the correct answer.\n"
        "   - Then, compare the student’s answer to the correct one in detail.\n"
        "   - If the answer is wrong, award **no more than 10%** of the marks for that criterion.\n"
        "   - Do not reward effort or partial understanding unless it's factually accurate and relevant.\n"
        f"   - Every criterion and the total score must be calculated out of {total_marks}.\n"
        "   - Your feedback must reflect this strict policy.\n"
        "3. Provide:\n"
        "   - Criterion-wise marks (with clear headings and marks out of the correct total)\n"
        "   - Detailed feedback for each criterion\n"
        f"   - A total score out of {total_marks} only (not 100)\n"
        "   - A summary comment\n\n"
        "Student Submission:\n"
        f"\"\"\"\n{submission_text}\n\"\"\"\n"
    )

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "meta-llama/llama-3.2-3b-instruct:free",
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    response = requests.post(API_URL, headers=headers, json=payload)
    response.raise_for_status()
    feedback_text = response.json()["choices"][0]["message"]["content"]

    # Save feedback to DOCX
    doc = Document()
    doc.add_heading("AI Teaching Assistant - Grading and Feedback", level=1)
    for line in feedback_text.strip().split("\n"):
        doc.add_paragraph(line.strip())

    output_dir = os.path.join("static", "feedback")
    os.makedirs(output_dir, exist_ok=True)
    doc.save(os.path.join(output_dir, "feedback.docx"))

    return feedback_text


def download_feedback():
    return send_from_directory(
        os.path.join("static", "feedback"),
        "feedback.docx",
        as_attachment=True
    )

